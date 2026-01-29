import urllib3
import warnings
import sys
import os
import re
import json
import time
import requests
import cv2
import csv
import numpy as np
import pdfplumber
import pytesseract
from pytesseract import Output
from PIL import Image

from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLineEdit, QLabel, QFileDialog, 
                             QTextEdit, QTreeWidget, QTreeWidgetItem, 
                             QSplitter, QProgressBar, QMessageBox)
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from PyQt5.QtGui import QColor

# Word 导出支持
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 禁用 SSL 安全警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
# 禁用 Python 弃用警告
warnings.filterwarnings("ignore", category=DeprecationWarning)

# --- 1. 配置区域 ---
TESSERACT_CMD = r'D:\Python\Scripts\tesseract.exe'
TESSDATA_DIR = r'D:\Python\Scripts\tessdata'

# 初始化 Tesseract
if os.path.exists(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
os.environ['TESSDATA_PREFIX'] = TESSDATA_DIR

# 局域网 LLM 配置
API_URL = "https://www.deepseek.com:18080/v1/chat/completions"
API_KEY = "your api key"

# --- 2. 图像预处理 ---
def preprocess_image_for_ocr(pil_image):
    """图像增强：灰度 -> 降噪 -> 二值化"""
    open_cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
    gray = cv2.medianBlur(gray, 3)
    _, binary_image = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return binary_image

# --- 3. 数据结构 ---
class DocumentLine:
    """存储每一行的结构化数据"""
    def __init__(self, text, role="BODY", page_num=0):
        self.text = text.strip()
        self.role = role
        self.page_num = page_num

# --- 4. 核心混合引擎 (Hybrid Engine) ---
class HybridPDFEngine:
    def __init__(self, dpi=300):
        self.dpi = dpi

    def log_console(self, msg):
        """
        工业级健壮的日志打印方法
        修复了 sys.stdout 为 None 时导致的 flush 崩溃问题
        """
        timestamp = time.strftime("%H:%M:%S")
        output_msg = f"[{timestamp}] {msg}"
        
        # 1. 优先尝试标准的 print，设置 flush=True (Python 3 内置支持)
        # 2. 增加 try-except 保护，防止在彻底没有 console 的环境下崩溃
        try:
            if sys.stdout is not None:
                print(output_msg, flush=True)
            else:
                # 如果没有标准输出流，我们至少保证它在 IDE 调试器里能看到
                import logging
                logging.info(output_msg)
        except (AttributeError, TypeError, OSError):
            # 最后的防线：如果上述都失败，直接忽略，确保不阻断业务逻辑
            pass

    def call_local_llm(self, messy_text, logger_callback=None):
        """
        调用局域网大模型，包含详细的交互日志记录
        """
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}"
        }
        
        system_prompt = (
            "你是一个OCR文档解析专家。请将以下破碎的文本修复为整洁的Markdown格式。"
            "要求：1. 识别并标记标题（# 一级, ## 二级）。"
            "2. 修复OCR导致的断句和明显错别字。"
            "3. 直接输出Markdown内容，不要任何开场白。"
        )
        
        payload = {
            "model": "DeepSeek-V3", 
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"破碎文本：\n\n{messy_text}"}
            ],
            "temperature": 0.3,
            "stream": False
        }

        # --- 交互日志：发送前 ---
        self.log_console(f"--- [LLM Request] ---\nPayload Size: {len(messy_text)} chars")
        if logger_callback:
            logger_callback(f"正在发送 {len(messy_text)} 字符到大模型...")

        try:
            start_time = time.time()
            response = requests.post(API_URL, headers=headers, json=payload, verify=False, timeout=120)
            elapsed = time.time() - start_time
            
            if response.status_code == 200:
                result = response.json()['choices'][0]['message']['content']
                # --- 交互日志：接收后 ---
                self.log_console(f"--- [LLM Response] ({elapsed:.2f}s) ---\nPreview: {result[:100]}...")
                return result
            else:
                err_msg = f"LLM Error: Status {response.status_code} - {response.text}"
                self.log_console(err_msg)
                return f"LLM 修复失败 (Code {response.status_code}):\n{messy_text}"
                
        except Exception as e:
            err_msg = f"局域网请求异常: {str(e)}"
            self.log_console(err_msg)
            return f"网络错误保留原稿:\n{messy_text}"

    def extract_and_merge_spatially(self, page_img):
        """基于物理坐标的初步行合并"""
        processed_img = preprocess_image_for_ocr(page_img)
        d = pytesseract.image_to_data(processed_img, lang='chi_sim', output_type=Output.DICT)
        
        lines = []
        current_line_text = []
        last_top = -1
        
        n_boxes = len(d['text'])
        for i in range(n_boxes):
            text = d['text'][i].strip()
            if not text: continue
            
            top = d['top'][i]
            height = d['height'][i]
            
            # 简单合并策略：垂直距离 < 行高一半
            if last_top == -1 or abs(top - last_top) < (height / 2):
                current_line_text.append(text)
            else:
                lines.append("".join(current_line_text))
                current_line_text = [text]
            last_top = top
            
        if current_line_text:
            lines.append("".join(current_line_text))
            
        return "\n".join(lines)

    def process_pdf(self, pdf_path, progress_callback=None, log_callback=None):
        results = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total = len(pdf.pages)
                self.log_console(f"开始处理文件: {pdf_path}, 共 {total} 页")
                
                for i, page in enumerate(pdf.pages):
                    p_idx = i + 1
                    msg = f"正在处理第 {p_idx}/{total} 页..."
                    if progress_callback: progress_callback(msg, int((p_idx/total)*100))
                    if log_callback: log_callback(msg)
                    
                    # 1. 物理提取 (OCR)
                    img = page.to_image(resolution=self.dpi).original
                    raw_text = self.extract_and_merge_spatially(img)
                    
                    # 2. 语义修复 (LLM)
                    markdown_text = self.call_local_llm(raw_text, log_callback)
                    
                    # 3. 结果解析
                    for line_text in markdown_text.split('\n'):
                        if not line_text.strip(): continue
                        role = "BODY"
                        if line_text.startswith('# '): role = "H1"
                        elif line_text.startswith('## '): role = "H2"
                        results.append(DocumentLine(line_text, role, p_idx))
                    
                    # 强制垃圾回收
                    del img
                    del raw_text
                    
        except Exception as e:
            self.log_console(f"Critical Error in process_pdf: {e}")
            raise e
            
        return results

# --- 5. UI 线程工作者 ---
class ParserWorker(QThread):
    progress_signal = pyqtSignal(str, int)
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(list) 
    error_signal = pyqtSignal(str)

    def __init__(self, filepath):
        super().__init__()
        self.filepath = filepath

    def run(self):
        try:
            engine = HybridPDFEngine()
            lines = engine.process_pdf(
                self.filepath, 
                progress_callback=self.progress_signal.emit,
                log_callback=self.log_signal.emit
            )
            self.finished_signal.emit(lines)
        except Exception as e:
            self.error_signal.emit(str(e))

# --- 6. 主界面 ---
class HybridRAGUI(QWidget):
    def __init__(self):
        super().__init__()
        self.all_parsed_lines = [] 
        self._init_ui()

    def _init_ui(self):
        self.setWindowTitle("Day 1: 视觉语义解析 (控制台透视版)")
        self.resize(1100, 750)
        
        main_layout = QVBoxLayout()
        
        # 顶部栏
        top_layout = QHBoxLayout()
        self.path_edit = QLineEdit()
        btn_file = QPushButton("浏览 PDF"); btn_file.clicked.connect(self.select_file)
        self.btn_run = QPushButton("🚀 开始解析"); self.btn_run.clicked.connect(self.run_engine)
        self.btn_run.setStyleSheet("background-color: #2874A6; color: white; font-weight: bold;")
        self.btn_export = QPushButton("💾 导出结果"); self.btn_export.clicked.connect(self.export_data)
        self.btn_export.setEnabled(False)
        
        top_layout.addWidget(QLabel("PDF:")); top_layout.addWidget(self.path_edit)
        top_layout.addWidget(btn_file); top_layout.addWidget(self.btn_run); top_layout.addWidget(self.btn_export)
        
        # 分割视图
        splitter = QSplitter(Qt.Horizontal)
        
        # 左侧：日志控制台
        log_widget = QWidget()
        log_layout = QVBoxLayout(log_widget)
        log_layout.addWidget(QLabel("交互日志 (LLM Interaction):"))
        self.log_box = QTextEdit()
        self.log_box.setReadOnly(True)
        self.log_box.setStyleSheet("background: #1E1E1E; color: #00FF00; font-family: Consolas; font-size: 10pt;")
        log_layout.addWidget(self.log_box)
        
        # 右侧：结构化树
        tree_widget = QWidget()
        tree_layout = QVBoxLayout(tree_widget)
        tree_layout.addWidget(QLabel("解析结果结构:"))
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["层级", "文本内容", "页码"])
        self.tree.setColumnWidth(0, 80); self.tree.setColumnWidth(2, 50)
        tree_layout.addWidget(self.tree)
        
        splitter.addWidget(log_widget)
        splitter.addWidget(tree_widget)
        splitter.setSizes([400, 700])
        
        # 底部进度
        self.pbar = QProgressBar()
        self.status_lbl = QLabel("就绪")
        
        main_layout.addLayout(top_layout)
        main_layout.addWidget(splitter)
        main_layout.addWidget(self.pbar)
        main_layout.addWidget(self.status_lbl)
        self.setLayout(main_layout)

    def select_file(self):
        f, _ = QFileDialog.getOpenFileName(self, "选择PDF", "", "PDF (*.pdf)")
        if f: self.path_edit.setText(f)

    def run_engine(self):
        path = self.path_edit.text()
        if not path or not os.path.exists(path):
            QMessageBox.warning(self, "错误", "路径无效")
            return
            
        self.btn_run.setEnabled(False); self.btn_export.setEnabled(False)
        self.tree.clear(); self.log_box.clear(); self.all_parsed_lines = []
        
        self.worker = ParserWorker(path)
        self.worker.progress_signal.connect(lambda m, v: (self.status_lbl.setText(m), self.pbar.setValue(v)))
        self.worker.log_signal.connect(self.append_log) 
        self.worker.finished_signal.connect(self.on_success)
        self.worker.error_signal.connect(self.on_fail)
        self.worker.start()

    def append_log(self, msg):
        self.log_box.append(f"> {msg}")
        self.log_box.verticalScrollBar().setValue(self.log_box.verticalScrollBar().maximum())

    def on_success(self, lines):
        self.btn_run.setEnabled(True); self.btn_export.setEnabled(True)
        self.all_parsed_lines = lines
        self.status_lbl.setText("解析完成")
        self.log_box.append("\n✅ 解析成功！可导出 JSON/Word/CSV。")
        
        for l in lines:
            item = QTreeWidgetItem([l.role, l.text, str(l.page_num)])
            if l.role == "H1":
                item.setBackground(0, QColor("#2980B9")); item.setForeground(0, QColor("white"))
            elif l.role == "H2":
                item.setBackground(0, QColor("#27AE60")); item.setForeground(0, QColor("white"))
            self.tree.addTopLevelItem(item)

    def on_fail(self, err):
        self.btn_run.setEnabled(True)
        QMessageBox.critical(self, "失败", f"引擎运行出错: {err}")
        self.log_box.append(f"\n❌ ERROR: {err}")

    def export_data(self):
        if not self.all_parsed_lines: return
        options = QFileDialog.Options()
        file_filter = "JSON Data (*.json);;Word (*.docx);;CSV (*.csv);;Text (*.txt)"
        fname, selected_filter = QFileDialog.getSaveFileName(self, "导出结果", "", file_filter, options=options)
        
        if not fname: return

        try:
            if "json" in selected_filter:
                data_export = [{"role": l.role, "text": l.text, "page": l.page_num} for l in self.all_parsed_lines]
                with open(fname, 'w', encoding='utf-8') as f:
                    json.dump(data_export, f, ensure_ascii=False, indent=2)
            elif "docx" in selected_filter:
                if not HAS_DOCX:
                    QMessageBox.warning(self, "警告", "未安装 python-docx 库")
                    return
                doc = Document()
                for l in self.all_parsed_lines:
                    if l.role == "H1": doc.add_heading(l.text.replace('# ', ''), 1)
                    elif l.role == "H2": doc.add_heading(l.text.replace('## ', ''), 2)
                    else: doc.add_paragraph(l.text)
                doc.save(fname)
            elif "csv" in selected_filter:
                with open(fname, 'w', encoding='utf-8-sig', newline='') as f:
                    writer = csv.writer(f)
                    writer.writerow(["层级", "内容", "页码"])
                    for l in self.all_parsed_lines:
                        writer.writerow([l.role, l.text, l.page_num])
            elif "txt" in selected_filter:
                with open(fname, 'w', encoding='utf-8') as f:
                    for l in self.all_parsed_lines: f.write(f"{l.text}\n")
            QMessageBox.information(self, "成功", f"文件导出成功: {fname}")
        except Exception as e:
            QMessageBox.critical(self, "导出错误", str(e))

if __name__ == "__main__":
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    app = QApplication(sys.argv)
    window = HybridRAGUI()
    window.show()
    sys.exit(app.exec_())